import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, color_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_heading_styled(doc, text, level):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    
    run = p.runs[0]
    run.font.name = 'Segoe UI'
    if level == 1:
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(16, 185, 129)  # Emerald green
        run.bold = True
    elif level == 2:
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(59, 130, 246)  # Royal Blue
        run.bold = True
    return p

def create_styled_table(doc, headers, data, col_widths=None):
    table = doc.add_table(rows=len(data) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    hdr_cells = table.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        set_cell_background(hdr_cells[i], "10B981")
        set_cell_margins(hdr_cells[i], top=120, bottom=120, left=150, right=150)
        p = hdr_cells[i].paragraphs[0]
        run = p.runs[0]
        run.font.name = 'Segoe UI'
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        
    for r_idx, row_data in enumerate(data):
        row_cells = table.rows[r_idx + 1].cells
        bg_color = "F8FAFC" if r_idx % 2 == 1 else "FFFFFF"
        for c_idx, val in enumerate(row_data):
            row_cells[c_idx].text = str(val)
            set_cell_background(row_cells[c_idx], bg_color)
            set_cell_margins(row_cells[c_idx], top=100, bottom=100, left=150, right=150)
            p = row_cells[c_idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            if len(p.runs) > 0:
                run = p.runs[0]
                run.font.name = 'Segoe UI'
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(51, 65, 85)
                
    if col_widths:
        for row in table.rows:
            for idx, width in enumerate(col_widths):
                row.cells[idx].width = Inches(width)
                
    doc.add_paragraph().paragraph_format.space_after = Pt(8)
    return table

def add_bullet_styled(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.left_indent = Inches(0.4)
    run = p.add_run(text)
    run.font.name = 'Segoe UI'
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(51, 65, 85)
    return p

def add_body_paragraph(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.font.name = 'Segoe UI'
        r_pre.font.size = Pt(11)
        r_pre.bold = True
        r_pre.font.color.rgb = RGBColor(15, 23, 42)
        
    run = p.add_run(text)
    run.font.name = 'Segoe UI'
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(51, 65, 85)
    return p

def add_callout(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    
    cell = table.rows[0].cells[0]
    set_cell_background(cell, "F1F5F9")
    set_cell_margins(cell, top=140, bottom=140, left=200, right=200)
    
    tcPr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        '<w:left w:val="single" w:sz="36" w:space="0" w:color="3B82F6"/>' # Blue callout border
        '<w:top w:val="none"/>'
        '<w:bottom w:val="none"/>'
        '<w:right w:val="none"/>'
        '</w:tcBorders>'
    )
    tcPr.append(borders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.font.name = 'Segoe UI'
    run.font.size = Pt(10.5)
    run.font.italic = True
    run.font.color.rgb = RGBColor(15, 23, 42)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(8)

def generate_demo_class_doc():
    doc = Document()
    
    # Page setup
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Header Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(20)
    title_p.paragraph_format.space_before = Pt(20)
    
    run = title_p.add_run("Antigravity Developer Demo Class\n")
    run.font.name = 'Segoe UI'
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(15, 23, 42)
    
    run_sub = title_p.add_run("Building Smart Prototypes from Scratch using AI-Agent Workflows")
    run_sub.font.name = 'Segoe UI'
    run_sub.font.size = Pt(14)
    run_sub.font.color.rgb = RGBColor(100, 116, 139)
    
    doc.add_page_break()

    # Introduction
    add_heading_styled(doc, "1. Workshop Introduction: What is Agentic Coding?", level=1)
    add_body_paragraph(doc, 
        "Welcome to the Antigravity Developer Demo Class. This guide outlines how to build a complete, highly complex mobile application "
        "from a clean slate without writing code manually. Agentic Coding is a paradigm shift where the human behaves as a Systems "
        "Architect, designing heuristics and requirements; Large Language Models (like GPT/Claude) act as Planners generating technical "
        "blueprints; and the Antigravity AI Agent acts as the Engineer executing code directories directly on the filesystem.")

    # Step 1
    add_heading_styled(doc, "Step 1: Install & Set Up the Antigravity IDE", level=2)
    add_body_paragraph(doc, 
        "Download the Antigravity IDE setup bundle. Set up an empty directory on your desktop or drive. Open the workspace in the IDE. "
        "The agent automatically scans your local system directories to discover active SDK paths (Flutter, Python, Android) to sync compilers.")

    # Step 2
    add_heading_styled(doc, "Step 2: Generate the Technical Blueprint with GPT/Claude", level=2)
    add_body_paragraph(doc, "Students craft a prompt explaining system needs (YOLO models, OpenCV, MethodChannels) to get a full plan:")
    add_callout(doc, "Sample Blueprint Prompt:\n\n'Act as an Expert AI & Mobile Architect. I want to build a Flutter application called Guard Cross AI that uses a phone camera and a local YOLOv8 TFLite model... Write a detailed implementation guide...'")

    # Step 3
    add_heading_styled(doc, "Step 3: Paste the Blueprint and Activate the Agent", level=2)
    add_body_paragraph(doc, 
        "Paste the blueprint into the Antigravity chat console. The agent parses requirements and starts a series of terminal steps: "
        "initializing Flutter, updating pubspec.yaml with dependencies, and running 'flutter pub get' to fetch packages in the background.")

    # Step 4
    add_heading_styled(doc, "Step 4: Proactive Asset & UI Synthesis", level=2)
    add_body_paragraph(doc, 
        "The agent automatically identifies missing graphic assets, calling its built-in generators to make app_logo.png and background_overlay.png. "
        "It then compiles frosted home dashboards and HUD radar screens featuring customized animated sweeping radar scan lines.")

    # Step 5
    add_heading_styled(doc, "Step 5: Dual-Mode and Native Integration", level=2)
    add_body_paragraph(doc, 
        "To allow running on Chrome (Web) or Emulators lacking hardware sensors, the agent structures a Simulation Sandbox console playing urban traffic "
        "scenarios. It configures Android permissions, build.gradle files, and Kotlin modules (preprocess, tracking, inference) automatically.")

    # Step 6
    add_heading_styled(doc, "Step 6: Automated Verification & Compile", level=2)
    add_body_paragraph(doc, 
        "The agent runs static analysis to remove warnings, injects custom SafetyEngine unit tests into widget_test.dart, "
        "executes tests, and initiates compile servers ('flutter run -d chrome') to open the working application.")

    # Section 3: Power of this Workflow (Table Format)
    add_heading_styled(doc, "2. Educational Power of Agentic Workflows", level=1)
    add_body_paragraph(doc, "Traditional development processes are compared to the agentic workflow in this validation matrix:")
    
    headers = ["Evaluation Metric", "Traditional Student Method", "Antigravity Agent Workflow"]
    data = [
        ["Package Mismatch", "Weeks spent resolving Kotlin/Gradle dependency compilation errors.", "Instant resolution as the agent updates configs dynamically."],
        ["Prototype Scope", "Simple static mockup screens with no real algorithmic logic.", "Full system architecture incorporating real IoU trackers & filters."],
        ["Presentation Risk", "High failure risk if live roadways are unavailable during demo.", "Simulation Sandbox allowing presenters to test scenarios on Chrome."]
    ]
    create_styled_table(doc, headers, data, [1.8, 2.5, 2.7])

    doc.save("Antigravity_Developer_Demo_Class.docx")
    print("Antigravity_Developer_Demo_Class.docx created successfully!")

if __name__ == "__main__":
    generate_demo_class_doc()
